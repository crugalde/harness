"""Diagnóstico previo a la primera corrida: qué falta para que el pipeline funcione.

Cada comprobación toca de verdad lo que va a usar el lote —escribe en la carpeta destino,
convierte un `.docx` real, lanza un chat de Hermes— en vez de suponer. Cuando algo falla,
dice el arreglo exacto en vez de un mensaje genérico.

    python hermes_brain.py comprobar
    python hermes_brain.py comprobar --rapido      # sin lanzar Hermes
"""
from __future__ import annotations

import importlib.util
import platform
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

from . import hermes as h
from .config import Config, ErrorConfig, cargar

OK, AVISO, FALLA = "ok", "aviso", "falla"
SIMBOLO = {OK: "  ok  ", AVISO: " aviso", FALLA: " FALTA"}
TIMEOUT_SONDA_S = 180


@dataclass
class Comprobacion:
    nombre: str
    estado: str
    detalle: str = ""
    arreglo: str = ""


@dataclass
class Diagnostico:
    items: list[Comprobacion] = field(default_factory=list)

    def añadir(self, nombre: str, estado: str, detalle: str = "", arreglo: str = "") -> Comprobacion:
        c = Comprobacion(nombre, estado, detalle, arreglo)
        self.items.append(c)
        return c

    @property
    def fallas(self) -> list[Comprobacion]:
        return [c for c in self.items if c.estado == FALLA]

    @property
    def avisos(self) -> list[Comprobacion]:
        return [c for c in self.items if c.estado == AVISO]


# --------------------------------------------------------------------------- comprobaciones
def _python(d: Diagnostico) -> None:
    v = sys.version_info
    if v >= (3, 10):
        d.añadir("Python", OK, f"{v.major}.{v.minor}.{v.micro} · {platform.system()}")
    else:
        d.añadir("Python", FALLA, f"{v.major}.{v.minor}: el worker necesita 3.10 o superior",
                 "Instala Python 3.10+ desde python.org y vuelve a crear el entorno.")


def _dependencias(d: Diagnostico, cfg: Config | None) -> None:
    necesarias = {"yaml": "pyyaml", "pypdf": "pypdf", "docx": "python-docx"}
    if cfg and cfg.n8n.base_url:
        necesarias["requests"] = "requests"
    faltan = [paquete for modulo, paquete in necesarias.items()
              if importlib.util.find_spec(modulo) is None]
    if faltan:
        d.añadir("Dependencias", FALLA, f"faltan: {', '.join(faltan)}",
                 f"pip install {' '.join(faltan)}")
    else:
        d.añadir("Dependencias", OK, ", ".join(necesarias.values()))


def _carpetas(d: Diagnostico, cfg: Config) -> None:
    from .inventario import recorrer, solo_en_la_nube

    if not cfg.carpetas:
        d.añadir("Carpeta a recorrer", OK,
                 "se indica en cada corrida (--carpeta), no en la configuración")
        return
    for carpeta in cfg.carpetas:
        if not carpeta.exists():
            d.añadir(f"Carpeta {carpeta}", FALLA, "no existe",
                     "Corrige la ruta en `carpetas:` del YAML.")
            continue
        vistos = nube = 0
        try:
            for ruta in recorrer(carpeta, cfg.extensiones, cfg.excluir):
                vistos += 1
                try:
                    if solo_en_la_nube(ruta.stat()):
                        nube += 1
                except OSError:
                    pass
                if vistos >= 2000:
                    break
        except PermissionError:
            d.añadir(f"Carpeta {carpeta}", FALLA, "sin permiso de lectura",
                     "Da acceso de lectura a tu usuario sobre esa carpeta.")
            continue
        cuenta = f"{vistos}+" if vistos >= 2000 else str(vistos)
        if vistos == 0:
            d.añadir(f"Carpeta {carpeta}", AVISO, f"sin archivos {'/'.join(cfg.extensiones)}",
                     "Revisa la ruta o amplía `extensiones:`.")
        elif nube and not cfg.procesar_solo_en_la_nube:
            d.añadir(f"Carpeta {carpeta}", AVISO,
                     f"{cuenta} archivos, de los cuales {nube}+ están solo en la nube y se saltarán",
                     "Clic derecho → «Conservar siempre en este dispositivo», o pon "
                     "`procesar_solo_en_la_nube: true`.")
        else:
            d.añadir(f"Carpeta {carpeta}", OK, f"{cuenta} archivos candidatos")


def _destino(d: Diagnostico, cfg: Config) -> None:
    try:
        cfg.destino_md.mkdir(parents=True, exist_ok=True)
        prueba = cfg.destino_md / ".hermes_brain_prueba"
        prueba.write_text("ok", encoding="utf-8")
        prueba.unlink()
    except OSError as exc:
        d.añadir("Destino brain md", FALLA, f"no se puede escribir: {exc}",
                 f"Comprueba que «{cfg.destino_md}» existe y que OneDrive no la tiene bloqueada.")
        return
    libre = shutil.disk_usage(cfg.destino_md).free // (1024 ** 3)
    estado = AVISO if libre < 5 else OK
    d.añadir("Destino brain md", estado, f"{cfg.destino_md} · {libre} GB libres",
             "Libera espacio antes de un lote grande." if estado == AVISO else "")


def _conversor(d: Diagnostico, cfg: Config) -> None:
    if not cfg.script_docx_md.exists():
        d.añadir("Conversor .docx → .md", FALLA, f"no existe: {cfg.script_docx_md}",
                 "Apunta `script_docx_md:` al archivo dentro del repo.")
        return
    if importlib.util.find_spec("docx") is None:
        d.añadir("Conversor .docx → .md", FALLA, "falta python-docx", "pip install python-docx")
        return
    import docx

    tmp = Path(tempfile.mkdtemp(prefix="hermes_comprobar_"))
    try:
        doc = docx.Document()
        doc.add_heading("Prueba de conversión", level=1)
        doc.add_heading("Diagnóstico", level=2)
        doc.add_paragraph("Documento generado por `comprobar` para validar el conversor.")
        entrada = tmp / "prueba.docx"
        doc.save(str(entrada))
        proc = subprocess.run(
            [cfg.python or sys.executable, str(cfg.script_docx_md), str(entrada),
             "--salida", str(tmp / "salida"), "--json"],
            capture_output=True, text=True, encoding="utf-8", errors="replace",
            timeout=120, check=False)
        if proc.returncode != 0 or not list((tmp / "salida").glob("*.md")):
            detalle = ((proc.stderr or proc.stdout) or "sin salida").strip()[-200:]
            d.añadir("Conversor .docx → .md", FALLA, detalle,
                     "Corre el script a mano para ver el error completo.")
        else:
            d.añadir("Conversor .docx → .md", OK, "convierte y escribe correctamente")
    except Exception as exc:   # cualquier fallo aquí es un 'falta', no un crash del diagnóstico
        d.añadir("Conversor .docx → .md", FALLA, f"{type(exc).__name__}: {exc}"[:200], "")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _ejecutable(d: Diagnostico, cfg: Config) -> bool:
    if not cfg.hermes.comando:
        d.añadir("Ejecutable de Hermes", FALLA, "`hermes.comando` está vacío",
                 "python hermes_brain.py detectar")
        return False
    exe = cfg.hermes.comando[0]
    ruta = Path(exe)
    if ruta.exists() or shutil.which(exe):
        d.añadir("Ejecutable de Hermes", OK, exe)
        return True
    d.añadir("Ejecutable de Hermes", FALLA, f"no se encuentra: {exe}",
             "python hermes_brain.py detectar")
    return False


def _sonda_hermes(cfg: Config, skill: str, palabra: str = "LISTO") -> tuple[bool, str]:
    """Lanza un chat mínimo y comprueba que la respuesta llega por stdout."""
    cmd = h._sustituir(cfg.hermes.comando, {
        "archivo": "", "skill": skill,
        "prompt": f"Responde únicamente con la palabra {palabra}, sin nada más.",
        "prompt_file": "", "destino": str(cfg.destino_md), "adjuntos": str(cfg.dir_adjuntos),
        "slug": "prueba", "salida_json": "",
    })
    try:
        codigo, salida = h._ejecutar(cmd, TIMEOUT_SONDA_S, cfg.hermes.cwd, cfg.hermes.env)
    except h.ErrorHermes as exc:
        return False, str(exc)
    except FileNotFoundError:
        return False, f"no se encuentra el ejecutable: {cmd[0]}"
    if codigo != 0:
        return False, f"código {codigo}: {salida.strip()[-200:]}"
    if palabra not in salida.upper():
        return False, f"responde, pero no como se le pidió: {salida.strip()[-160:]}"
    return True, salida.strip()[-160:]


def _hermes(d: Diagnostico, cfg: Config) -> None:
    ok, detalle = _sonda_hermes(cfg, skill="")
    if not ok:
        d.añadir("Hermes responde", FALLA, detalle,
                 "Ejecuta el comando a mano en PowerShell para ver el error completo.")
        return
    d.añadir("Hermes responde", OK, "un chat de un disparo devuelve texto por stdout")

    for etiqueta, skill in (("PDF", cfg.hermes.skill_pdf), ("Word", cfg.hermes.skill_docx)):
        if not skill:
            d.añadir(f"Skill {etiqueta}", AVISO, "sin nombre configurado",
                     "El chat se abrirá sin precargar skill; el prompt se basta solo.")
            continue
        if etiqueta == "Word" and not cfg.hermes.revisar_docx_con_hermes:
            d.añadir(f"Skill {etiqueta}", OK, "no se usa: la revisión con Hermes está apagada")
            continue
        ok, detalle = _sonda_hermes(cfg, skill=skill)
        if ok:
            d.añadir(f"Skill {etiqueta} «{skill}»", OK, "se precarga sin error")
        else:
            d.añadir(f"Skill {etiqueta} «{skill}»", FALLA, detalle,
                     "Confirma el nombre exacto de la skill y corrígelo en el YAML.")


def _n8n(d: Diagnostico, cfg: Config) -> None:
    if not cfg.n8n.base_url:
        d.añadir("n8n (VPS)", OK, "sin configurar: el worker trabaja igual, solo sin panel")
        return
    from .cliente_n8n import ClienteN8n

    cliente = ClienteN8n(cfg.n8n)
    control = cliente.control("comprobacion")
    if cliente.ultimo_error:
        d.añadir("n8n (VPS)", FALLA, cliente.ultimo_error[:200],
                 "Revisa que el flujo esté ACTIVO, la URL termine en /webhook y que "
                 "HERMES_TOKEN sea el mismo en el contenedor y en el YAML.")
    else:
        d.añadir("n8n (VPS)", OK, f"responde y acepta el token (acción: {control.accion})")


# --------------------------------------------------------------------------- orquestación
def diagnosticar(ruta_config: str | None = None, rapido: bool = False) -> Diagnostico:
    d = Diagnostico()
    _python(d)
    try:
        cfg = cargar(ruta_config)
    except ErrorConfig as exc:
        d.añadir("Configuración", FALLA, str(exc).splitlines()[0],
                 "Copia tools/hermes_brain/config.example.yaml a "
                 "~/.config/harness/hermes_brain.yaml y ajústalo.")
        _dependencias(d, None)
        return d
    carpetas = f"{len(cfg.carpetas)} carpeta(s) por defecto" if cfg.carpetas else "sin carpeta fija"
    d.añadir("Configuración", OK, f"{carpetas} · destino {cfg.destino_md.name}")
    _dependencias(d, cfg)
    _carpetas(d, cfg)
    _destino(d, cfg)
    _conversor(d, cfg)
    if _ejecutable(d, cfg) and not rapido:
        _hermes(d, cfg)
    elif rapido:
        d.añadir("Hermes responde", AVISO, "no se probó (--rapido)",
                 "Quita --rapido para lanzar un chat de verdad.")
    _n8n(d, cfg)
    return d


ANCHO_NOMBRE = 34


def _acortar(texto: str, tope: int = ANCHO_NOMBRE) -> str:
    """Recorta por el medio: en una ruta larga, lo informativo son los dos extremos."""
    if len(texto) <= tope:
        return texto
    mitad = (tope - 1) // 2
    return f"{texto[:mitad]}…{texto[-(tope - mitad - 1):]}"


def formatear(d: Diagnostico) -> str:
    ancho = min(max((len(c.nombre) for c in d.items), default=20), ANCHO_NOMBRE)
    lineas = ["Diagnóstico del pipeline documental", "=" * 70, ""]
    for c in d.items:
        lineas.append(f"[{SIMBOLO[c.estado]}]  {_acortar(c.nombre).ljust(ancho)}  {c.detalle}")
        if c.arreglo:
            lineas.append(f"{' ' * (ancho + 12)}→ {c.arreglo}")
    lineas.append("")
    if d.fallas:
        lineas.append(f"Falta resolver {len(d.fallas)} punto(s) antes de correr un lote:")
        lineas += [f"  · {_acortar(c.nombre)}: {c.arreglo or c.detalle}" for c in d.fallas]
    elif d.avisos:
        lineas.append(f"Todo lo necesario está listo. {len(d.avisos)} aviso(s) que conviene mirar.")
        lineas.append("Siguiente paso: python hermes_brain.py probar-hermes \"<un PDF>\"")
    else:
        lineas.append("Todo listo. Siguiente paso:")
        lineas.append('  python hermes_brain.py probar-hermes "C:\\ruta\\a\\un_paper.pdf"')
    return "\n".join(lineas)
