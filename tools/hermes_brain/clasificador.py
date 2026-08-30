"""Reconocimiento de documentos: PDF de revista científica y Word de resumen clínico.

La clasificación es determinista y explicable: cada señal aporta un peso y queda registrada
en `evidencia`, de modo que un archivo dudoso pueda revisarse después sabiendo *por qué*
quedó dudoso. Nada de esto sale del PC (R8).
"""
from __future__ import annotations

import logging
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

# --------------------------------------------------------------------------- patrones
RE_DOI = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Za-z0-9]{4,}", re.I)
RE_ISSN = re.compile(r"\b(?:ISSN|eISSN)[:\s]*\d{4}-?\d{3}[\dxX]\b|\b\d{4}-\d{3}[\dxX]\b")
RE_ABSTRACT = re.compile(r"^\s*(a\s*b\s*s\s*t\s*r\s*a\s*c\s*t|abstract|resumen|summary)\b[\s:.-]*$",
                         re.I | re.M)
RE_ABSTRACT_INLINE = re.compile(r"\b(abstract|resumen)\s*[:—-]\s*\S", re.I)
RE_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]{2,}")
RE_RECIBIDO = re.compile(r"\b(received|accepted|revised|recibido|aceptado)\b.{0,40}\b(20\d\d|19\d\d)", re.I)
RE_KEYWORDS = re.compile(r"^\s*(key\s?words|palabras\s+clave)\b", re.I | re.M)
RE_REFERENCIAS = re.compile(r"^\s*(references|referencias|bibliography|bibliograf[ií]a)\s*$", re.I | re.M)
RE_PMID = re.compile(r"\bPMID:?\s*\d{6,9}\b", re.I)
RE_AUTORES = re.compile(
    r"^[A-ZÁÉÍÓÚÑ][\w.'’-]+(?:\s+[A-ZÁÉÍÓÚÑ][\w.'’-]+){0,3}"
    r"(?:\s*[a-z]?\s*[,;]\s*[A-ZÁÉÍÓÚÑ][\w.'’-]+(?:\s+[A-ZÁÉÍÓÚÑ][\w.'’-]+){0,3}){1,}", re.M)

EDITORIALES = ("elsevier", "springer", "wiley", "lippincott", "sage publications", "taylor & francis",
               "oxford university press", "bmj", "jama network", "nature publishing", "frontiers media",
               "mdpi", "karger", "thieme", "american academy of neurology", "sciencedirect",
               "pubmed", "medline", "j o u r n a l")
PISTAS_JOURNAL = ("journal of", "revista ", "annals of", "archives of", "the lancet", "new england journal",
                  "neurology", "brain ", "muscle & nerve", "clinical neurophysiology", "vol.", "volume ",
                  "issue ", "no. ", "©", "doi.org", "published by", "publicado por")
AFILIACION = ("university", "universidad", "department", "departamento", "hospital", "clinic", "clínica",
              "institute", "instituto", "school of medicine", "facultad", "servicio de", "unit ", "unidad de")

SECCIONES_PATOLOGIA = {
    "definicion": ("definición", "definicion", "concepto", "introducción", "introduccion", "generalidades"),
    "epidemiologia": ("epidemiología", "epidemiologia", "incidencia", "prevalencia"),
    "etiologia": ("etiología", "etiologia", "causas", "factores de riesgo"),
    "fisiopatologia": ("fisiopatología", "fisiopatologia", "patogenia", "patogénesis", "patogenesis"),
    "clinica": ("clínica", "clinica", "cuadro clínico", "manifestaciones clínicas", "síntomas", "sintomas",
                "presentación clínica", "signos y síntomas"),
    "diagnostico": ("diagnóstico", "diagnostico", "criterios diagnósticos", "estudio diagnóstico",
                    "exámenes complementarios", "examenes complementarios", "laboratorio", "imagenología"),
    "diferencial": ("diagnóstico diferencial", "diagnostico diferencial", "diferenciales"),
    "tratamiento": ("tratamiento", "manejo", "terapia", "terapéutica", "terapeutica", "manejo inicial"),
    "pronostico": ("pronóstico", "pronostico", "evolución y pronóstico", "complicaciones", "seguimiento"),
    "clasificacion_clinica": ("clasificación", "clasificacion", "tipos", "formas clínicas", "estadificación"),
}
MARCADORES_FICHA = ("motivo de consulta", "anamnesis próxima", "anamnesis proxima", "anamnesis remota",
                    "examen físico", "examen fisico", "evolución clínica", "hospitalizado en",
                    "ficha clínica", "ficha clinica", "n° ficha", "rut:", "r.u.t", "paciente de ",
                    "ingresa el ", "epicrisis", "identificación del paciente")
VOCABULARIO_CLINICO = ("paciente", "síntomas", "sintomas", "signos", "diagnóstico", "diagnostico",
                       "tratamiento", "dosis", "mg/día", "mg/kg", "sensibilidad", "especificidad",
                       "pronóstico", "biopsia", "terapia", "clínico", "clinico", "patología", "patologia",
                       "enfermedad", "síndrome", "sindrome", "corticoides", "inmunosupresor")
ANTI_CLINICO = ("factura", "boleta", "cotización", "cotizacion", "orden de compra", "curriculum vitae",
                "curriculum", "acta de reunión", "acta de reunion", "contrato de", "carta de renuncia",
                "programa del curso", "lista de asistencia", "presupuesto", "declaración de impuestos")

MAX_TEXTO = 40_000

# pypdf avisa por stderr de cada PDF malformado ("EOF marker not found"); eso ya se refleja
# en la clasificación y sin esto rompe la línea de avance del lote.
logging.getLogger("pypdf").setLevel(logging.ERROR)


@dataclass
class Clasificacion:
    """Veredicto de un archivo, con las señales que lo sostienen."""

    decision: str                      # cientifico | no_cientifico | clinico | no_clinico | dudoso | error
    score: float = 0.0
    evidencia: dict = field(default_factory=dict)
    motivo: str = ""
    titulo: str = ""

    @property
    def procesable(self) -> bool:
        return self.decision in ("cientifico", "clinico")


# --------------------------------------------------------------------------- extracción
def texto_pdf(ruta: Path, paginas: int = 2) -> tuple[str, str, dict]:
    """Devuelve (texto de las primeras páginas, texto de la última, metadatos)."""
    from pypdf import PdfReader

    lector = PdfReader(str(ruta), strict=False)
    if getattr(lector, "is_encrypted", False):
        try:
            lector.decrypt("")
        except Exception as exc:   # pypdf lanza varias clases según el tipo de cifrado
            raise ValueError(f"PDF cifrado: {exc}") from exc
    meta = {}
    try:
        for clave, valor in (lector.metadata or {}).items():
            meta[str(clave).lstrip("/").lower()] = str(valor)
    except Exception:          # metadatos corruptos: el PDF sigue siendo legible
        meta = {}
    partes = []
    for pagina in lector.pages[:paginas]:
        try:
            partes.append(pagina.extract_text() or "")
        except Exception:      # una página ilegible no invalida el resto del documento
            partes.append("")
    ultima = ""
    if len(lector.pages) > paginas:
        try:
            ultima = lector.pages[-1].extract_text() or ""
        except Exception:      # la última página solo aporta la señal de "referencias"
            ultima = ""
    meta["n_paginas"] = str(len(lector.pages))
    return "\n".join(partes)[:MAX_TEXTO], ultima[:MAX_TEXTO], meta


def texto_docx(ruta: Path, max_palabras: int = 1200) -> tuple[str, list[str]]:
    """Devuelve (texto plano acotado, lista de títulos de estilo Heading)."""
    import docx

    doc = docx.Document(str(ruta))
    lineas: list[str] = []
    titulos: list[str] = []
    palabras = 0
    for parrafo in doc.paragraphs:
        txt = parrafo.text.strip()
        if not txt:
            continue
        estilo = (parrafo.style.name or "").lower() if parrafo.style is not None else ""
        if estilo.startswith(("heading", "título", "titulo")) or (
            len(txt) < 90 and parrafo.runs and all(r.bold for r in parrafo.runs if r.text.strip())
        ):
            titulos.append(txt)
        lineas.append(txt)
        palabras += len(txt.split())
        if palabras >= max_palabras:
            break
    for tabla in doc.tables[:5]:
        for fila in tabla.rows[:20]:
            lineas.append(" | ".join(c.text.strip() for c in fila.cells))
    return "\n".join(lineas)[:MAX_TEXTO], titulos


def convertir_doc_a_docx(ruta: Path, soffice: str = "soffice", timeout_s: int = 180) -> Path:
    """Convierte un .doc legado a .docx con LibreOffice y devuelve la ruta temporal."""
    salida = Path(tempfile.mkdtemp(prefix="hermes_doc_"))
    proc = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(salida), str(ruta)],
        capture_output=True, text=True, timeout=timeout_s, check=False)
    convertido = salida / (ruta.stem + ".docx")
    if not convertido.exists():
        raise RuntimeError(f"LibreOffice no pudo convertir {ruta.name}: {proc.stderr.strip()[:200]}")
    return convertido


# --------------------------------------------------------------------------- PDF
def clasificar_pdf(ruta: Path, umbral_si: float = 4.0, umbral_no: float = 2.0,
                   paginas: int = 2) -> Clasificacion:
    """¿Es un PDF de revista científica? Exige título, journal, autor y abstract."""
    try:
        cabeza, cola_txt, meta = texto_pdf(ruta, paginas)
    except Exception as exc:   # un PDF corrupto es un resultado del lote, no un crash
        return Clasificacion("error", 0.0, {"excepcion": type(exc).__name__}, f"PDF ilegible: {exc}")

    if len(cabeza.strip()) < 120:
        return Clasificacion("dudoso", 0.0, {"sin_texto": True, "n_paginas": meta.get("n_paginas", "?")},
                             "PDF sin capa de texto (probable escaneo): requiere OCR o revisión manual")

    bajo = cabeza.lower()
    lineas = [l.strip() for l in cabeza.splitlines() if l.strip()]
    ev: dict = {}
    score = 0.0

    doi = RE_DOI.search(cabeza) or RE_DOI.search(meta.get("subject", "") + meta.get("keywords", ""))
    ev["doi"] = bool(doi)
    if doi:
        score += 2.0

    hay_abstract = bool(RE_ABSTRACT.search(cabeza) or RE_ABSTRACT_INLINE.search(cabeza))
    ev["abstract"] = hay_abstract
    if hay_abstract:
        score += 1.5

    pistas_j = [p for p in PISTAS_JOURNAL if p in bajo] + [e for e in EDITORIALES if e in bajo]
    hay_issn = bool(RE_ISSN.search(cabeza))
    hay_journal = bool(doi) or hay_issn or len(pistas_j) >= 2
    ev["journal"] = hay_journal
    ev["pistas_journal"] = pistas_j[:6]
    ev["issn"] = hay_issn
    if hay_journal:
        score += 1.5

    afiliaciones = [a for a in AFILIACION if a in bajo]
    autor_meta = _autor_creible(meta.get("author", ""))
    autores_linea = bool(RE_AUTORES.search("\n".join(lineas[:12])))
    hay_autor = autor_meta or (autores_linea and bool(afiliaciones or RE_EMAIL.search(cabeza)))
    ev["autor"] = hay_autor
    ev["afiliaciones"] = afiliaciones[:4]
    if hay_autor:
        score += 1.5

    titulo = _titulo_pdf(meta, lineas)
    ev["titulo"] = bool(titulo)
    if titulo:
        score += 1.0

    if RE_RECIBIDO.search(cabeza):
        ev["fechas_editoriales"] = True
        score += 0.5
    if RE_KEYWORDS.search(cabeza):
        ev["keywords"] = True
        score += 0.5
    if RE_REFERENCIAS.search(cola_txt) or RE_PMID.search(cola_txt):
        ev["referencias"] = True
        score += 0.5

    nucleo = sum(1 for k in ("titulo", "journal", "autor", "abstract") if ev.get(k))
    ev["nucleo"] = nucleo
    ev["n_paginas"] = meta.get("n_paginas", "?")

    if nucleo == 4 or (nucleo == 3 and ev["doi"]):
        decision, motivo = "cientifico", f"{nucleo}/4 elementos (título, journal, autor, abstract)"
    elif nucleo >= 3 and score >= umbral_si:
        decision, motivo = "cientifico", f"{nucleo}/4 elementos y score {score:.1f}"
    elif score <= umbral_no or nucleo <= 1:
        decision, motivo = "no_cientifico", f"solo {nucleo}/4 elementos (score {score:.1f})"
    else:
        decision, motivo = "dudoso", f"{nucleo}/4 elementos, score {score:.1f} en zona gris"
    return Clasificacion(decision, round(score, 2), ev, motivo, titulo)


AUTORES_GENERICOS = ("anonymous", "unknown", "user", "usuario", "admin", "administrador",
                     "windows user", "pc", "owner", "propietario", "acrobat", "word", "none")


def _autor_creible(bruto: str) -> bool:
    """Word y las impresoras a PDF rellenan /Author con el nombre de la cuenta del PC.

    Solo cuenta como autoría de artículo una lista de personas: coma, punto y coma, "and"
    o "&" separando nombres, o al menos dos palabras que no sean un genérico conocido.
    """
    autor = (bruto or "").strip()
    if len(autor) < 5 or autor.lower() in AUTORES_GENERICOS:
        return False
    if any(g == autor.lower() or g in autor.lower().split() for g in AUTORES_GENERICOS):
        return False
    if any(sep in autor for sep in (",", ";", " and ", " & ")):
        return True
    return len(autor.split()) >= 2


def _titulo_pdf(meta: dict, lineas: list[str]) -> str:
    """Título del artículo: metadato si es creíble, si no la primera línea con forma de título."""
    bruto = (meta.get("title") or "").strip()
    basura = ("untitled", "microsoft word", "document", "pdf", "print", "layout", ".doc", ".qxd", ".indd")
    if 15 <= len(bruto) <= 300 and not any(b in bruto.lower() for b in basura):
        return bruto
    for linea in lineas[:8]:
        if (20 <= len(linea) <= 250 and len(linea.split()) >= 4 and not RE_EMAIL.search(linea)
                and sum(c.isdigit() for c in linea) / max(len(linea), 1) < 0.2):
            return linea
    return ""


# --------------------------------------------------------------------------- Word
def clasificar_docx(ruta: Path, umbral_si: float = 4.0, umbral_no: float = 2.0,
                    max_palabras: int = 1200) -> Clasificacion:
    """¿Es un Word de resumen clínico de patologías? Ante la duda devuelve 'dudoso'."""
    try:
        texto, titulos = texto_docx(ruta, max_palabras)
    except Exception as exc:   # ídem: se registra como error y el lote continúa
        return Clasificacion("error", 0.0, {"excepcion": type(exc).__name__}, f"Word ilegible: {exc}")

    if len(texto.strip()) < 200:
        return Clasificacion("dudoso", 0.0, {"sin_texto": True},
                             "Documento casi vacío o solo con imágenes: requiere revisión")

    bajo = texto.lower()
    bajo_titulos = " \n".join(titulos).lower()
    ev: dict = {}
    score = 0.0

    secciones = [nombre for nombre, alias in SECCIONES_PATOLOGIA.items()
                 if any(a in bajo_titulos for a in alias)]
    secciones_texto = [nombre for nombre, alias in SECCIONES_PATOLOGIA.items()
                       if nombre not in secciones and any(a in bajo for a in alias)]
    ev["secciones_en_titulos"] = secciones
    ev["secciones_en_texto"] = secciones_texto
    score += min(len(secciones) * 1.0, 4.0) + min(len(secciones_texto) * 0.35, 1.5)

    marcadores = [m for m in MARCADORES_FICHA if m in bajo]
    ev["marcadores_ficha"] = marcadores[:6]
    ev["phi_probable"] = bool(marcadores)
    if marcadores:
        score += 1.5

    vocab = sorted({v for v in VOCABULARIO_CLINICO if v in bajo})
    ev["vocabulario"] = vocab[:10]
    score += min(len(vocab) * 0.2, 1.5)

    anti = [a for a in ANTI_CLINICO if a in bajo]
    ev["anti_clinico"] = anti
    score -= 2.0 * len(anti)

    ev["n_titulos"] = len(titulos)
    if len(titulos) >= 3:
        score += 0.5

    if anti and not secciones:
        decision, motivo = "no_clinico", f"marcadores no clínicos: {', '.join(anti[:3])}"
    elif score >= umbral_si:
        decision = "clinico"
        motivo = (f"{len(secciones)} secciones de patología en títulos "
                  f"({', '.join(secciones[:4])}), score {score:.1f}")
    elif score <= umbral_no:
        decision, motivo = "no_clinico", f"sin estructura clínica reconocible (score {score:.1f})"
    else:
        decision = "dudoso"
        motivo = (f"score {score:.1f} en zona gris; secciones: {', '.join(secciones) or 'ninguna'}; "
                  f"vocabulario clínico: {len(vocab)} términos")
    titulo = titulos[0] if titulos and not ev["phi_probable"] else ""
    return Clasificacion(decision, round(score, 2), ev, motivo, titulo)


def clasificar(ruta: Path, cfg) -> Clasificacion:
    """Despacha según extensión. `.doc` se convierte con LibreOffice si está habilitado."""
    ext = ruta.suffix.lower()
    if ext == ".pdf":
        return clasificar_pdf(ruta, cfg.pdf_umbral_si, cfg.pdf_umbral_no, cfg.paginas_pdf)
    if ext == ".docx":
        return clasificar_docx(ruta, cfg.docx_umbral_si, cfg.docx_umbral_no, cfg.palabras_docx)
    if ext == ".doc":
        if not cfg.convertir_doc_con_soffice:
            return Clasificacion("dudoso", 0.0, {"formato_legado": True},
                                 ".doc legado: habilita 'convertir_doc_con_soffice' o guárdalo como .docx")
        try:
            convertido = convertir_doc_a_docx(ruta, cfg.soffice)
        except Exception as exc:   # LibreOffice falla de muchas formas; todas dan 'error'
            return Clasificacion("error", 0.0, {"excepcion": type(exc).__name__}, str(exc)[:200])
        clf = clasificar_docx(convertido, cfg.docx_umbral_si, cfg.docx_umbral_no, cfg.palabras_docx)
        clf.evidencia["convertido_desde_doc"] = str(convertido)
        return clf
    return Clasificacion("no_soportado", 0.0, {"ext": ext}, f"extensión no soportada: {ext}")
