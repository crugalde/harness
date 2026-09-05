#!/usr/bin/env python3
"""paper_review.py — Análisis científico de múltiples papers (PDF/DOCX) con modelo por etapa.

Pipeline de 6 etapas, cada una con el tier de modelo que le corresponde (`model_policy`):

  1. descubrir   carpeta -> lista de PDF/DOCX/MD                        (sin modelo)
  2. extraer     texto crudo con pypdf / python-docx                    (sin modelo)
  3. de-identificar  RUT, nombres, fichas, contactos, fechas de nacimiento  (sin modelo, R8)
  4. fichar      por paper: diseño, n, población, hallazgo, límites     (T2 · sonnet-5)
  5. contrastar  PubMed por paper -> qué aporta frente a lo publicado   (tool + T2)
  6. sintetizar  lectura transversal a nivel de neurólogo académico     (T3 · opus-5)

Por qué así: fichar N papers son N llamadas y ahí manda la relación capacidad/velocidad
(Sonnet 5); la lectura transversal es UNA llamada y es donde se gana o se pierde el
análisis, así que ahí se paga capacidad máxima (Opus 5). Pagar Opus por paper multiplica
el costo sin mejorar la ficha; pagar Sonnet la síntesis final abarata lo único que no
conviene abaratar.

Salidas en `--out`:
  revision.json  — un registro estructurado por paper + la síntesis (encadenable)
  revision.md    — informe pegable en Notion, con la tabla comparativa y los PMIDs

Reglas que respeta: R2 (solo PMIDs devueltos por la tool), R8 (de-identifica antes de
cualquier llamada externa), R12 (declara truncados, fallos y confianza).

Uso:
  python tools/paper_review.py --dir ~/papers --tema "HD-sEMG en ELA" \
      --out projects/2026-09-01_hdsemg
  python tools/paper_review.py --dir ~/papers --dry-run      # extracción + de-ID, sin API
Req: pip install pypdf python-docx  (y biopython para el contraste PubMed)
"""
from __future__ import annotations

import argparse
import importlib.util
import json
import re
import sys
from dataclasses import dataclass, field, asdict
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(Path(__file__).resolve().parent))

try:
    from env_loader import load_env
    load_env()
except Exception:  # el harness debe correr aunque no haya .env local
    pass

import model_policy as mp      # noqa: E402
import backends                # noqa: E402

DOC_SUFFIXES = {".pdf", ".docx", ".md", ".txt"}
PAPER_CHAR_BUDGET = 60_000     # ~15k tokens por paper; el resto se declara truncado


# ---------------------------------------------------------------------------
# 1-2 · Descubrir y extraer
# ---------------------------------------------------------------------------
def discover(folder: Path) -> list[Path]:
    if not folder.is_dir():
        raise FileNotFoundError(f"No existe la carpeta: {folder}")
    return sorted(p for p in folder.rglob("*")
                  if p.is_file() and p.suffix.lower() in DOC_SUFFIXES)


def extract_text(path: Path) -> str:
    """Texto plano del documento. Falla con mensaje claro, nunca en silencio."""
    suf = path.suffix.lower()
    if suf in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")
    if suf == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise RuntimeError("Falta pypdf para leer PDFs: pip install pypdf") from e
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if suf == ".docx":
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError("Falta python-docx para leer .docx: pip install python-docx") from e
        doc = Document(str(path))
        partes = [p.text for p in doc.paragraphs]
        for tabla in doc.tables:
            for fila in tabla.rows:
                partes.append(" | ".join(c.text for c in fila.cells))
        return "\n".join(partes)
    raise RuntimeError(f"Formato no soportado: {path.suffix}")


# ---------------------------------------------------------------------------
# 3 · De-identificación (R8) — se aplica ANTES de cualquier salida externa
# ---------------------------------------------------------------------------
PHI_PATTERNS: list[tuple[str, str]] = [
    (r"\b\d{1,2}\.\d{3}\.\d{3}[-‐-]?[\dkK]\b", "[RUT]"),
    (r"\b\d{7,8}[-‐-][\dkK]\b", "[RUT]"),
    (r"[\w.\-+]+@[\w\-]+\.[\w.\-]+", "[EMAIL]"),
    (r"(?:\+?56)?\s?9\s?\d{4}\s?\d{4}\b", "[TELEFONO]"),
    (r"(?i)\b(?:paciente|nombre|apellidos?)\s*:\s*[^\n,;]{3,60}", "[NOMBRE]"),
    (r"(?i)\b(?:ficha|hc|historia\s+cl[ií]nica|mrn)\s*(?:n[°ºo]\.?)?\s*:?\s*\d{3,10}\b",
     "[FICHA]"),
    (r"(?i)\b(?:fecha\s+de\s+nacimiento|f\.?\s*nac\.?|fn)\s*:?\s*\d{1,2}[/-]\d{1,2}[/-]\d{2,4}",
     "[FNAC]"),
]


def deidentify(text: str) -> tuple[str, int]:
    """Devuelve (texto de-identificado, nº de redacciones). Conservador por diseño."""
    total = 0
    for pattern, tag in PHI_PATTERNS:
        text, n = re.subn(pattern, tag, text)
        total += n
    return text, total


# ---------------------------------------------------------------------------
# Registro por paper
# ---------------------------------------------------------------------------
@dataclass
class PaperRecord:
    archivo: str
    formato: str
    caracteres: int
    truncado: bool = False
    redacciones_phi: int = 0
    ficha: dict = field(default_factory=dict)
    pmids_relacionados: list[str] = field(default_factory=list)
    aporte: str = ""
    confianza: str = "no evaluada"
    errores: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Utilidades de modelo
# ---------------------------------------------------------------------------
def _json_from(text: str) -> dict:
    """Extrae el primer objeto JSON de la respuesta. Sin regex frágil sobre el texto."""
    depth, start = 0, None
    for i, ch in enumerate(text):
        if ch == "{":
            if depth == 0:
                start = i
            depth += 1
        elif ch == "}" and depth:
            depth -= 1
            if depth == 0 and start is not None:
                try:
                    return json.loads(text[start:i + 1])
                except json.JSONDecodeError:
                    start = None
    raise ValueError("la respuesta no contenía un objeto JSON válido")


FICHA_PROMPT = """Eres un neurólogo académico fichando un artículo para una revisión.
Devuelve SOLO un objeto JSON, sin texto alrededor, con estas claves exactas:

  titulo, autores, anio, revista, doi, diseno, n, poblacion, intervencion, comparador,
  desenlaces, hallazgo_principal, limitaciones, calidad_metodologica, terminos_pubmed

Reglas:
- `terminos_pubmed`: lista de 2 a 4 consultas MeSH/texto libre para buscar la literatura
  con la que este trabajo debe compararse. En inglés.
- `n`: número de sujetos como entero, o null si no se declara.
- `calidad_metodologica`: una frase con el riesgo de sesgo principal.
- Si un dato NO está en el texto, pon null. NO lo infieras ni lo completes de memoria.

TEMA DE LA REVISIÓN: {tema}
ARCHIVO: {archivo}

TEXTO DEL ARTÍCULO (de-identificado):
---
{texto}
---"""

APORTE_PROMPT = """Eres un neurólogo académico. Tienes la ficha de un artículo y los
abstracts que PubMed devolvió para su literatura de referencia.

Determina QUÉ APORTA este artículo frente a lo ya publicado. Devuelve SOLO un JSON con:

  aporte            — 2 a 4 frases: qué añade (o confirma, o contradice) y por qué importa.
  novedad           — "nuevo" | "confirmatorio" | "contradictorio" | "incremental" |
                      "no determinable"
  pmids_citados     — lista de PMIDs que SÍ aparecen en los resultados de abajo. Si no hay
                      resultados, lista vacía. NUNCA escribas un PMID que no esté abajo.
  confianza         — "alta" | "media" | "baja", según cuánto texto real respalde el juicio.
  no_verificado     — lista de afirmaciones que no pudiste anclar a los textos disponibles.

FICHA DEL ARTÍCULO:
{ficha}

RESULTADOS DE PUBMED (única fuente válida de PMIDs):
---
{pubmed}
---"""

SINTESIS_PROMPT = """Eres un neurólogo académico escribiendo la lectura transversal de un
conjunto de artículos para un par, no para un estudiante. Tono directo, sin relleno.

Escribe en español, en Markdown, con estas secciones:

## Pregunta y alcance
## Qué muestra el conjunto
Convergencias reales entre los trabajos, con el peso metodológico de cada uno.
## Dónde se contradicen
Discrepancias y la explicación más probable (población, técnica, definición de desenlace).
## Calidad de la evidencia
Riesgo de sesgo dominante del conjunto y qué diseño haría falta para resolverlo.
## Aporte neto
Qué sabemos ahora que no sabíamos, en una lista de puntos.
## Vacíos y qué haría falta
Preguntas abiertas, en orden de importancia clínica.
## Límites de esta revisión
Qué quedó fuera, qué no se pudo verificar y con qué confianza.

Reglas duras:
- Cita únicamente los PMIDs presentes en las fichas. Si un dato no está respaldado,
  márcalo `[sin verificar]` en línea; no lo presentes como establecido.
- No nombres un gen como causa de una enfermedad si no aparece en el material.
- La ausencia de dato es un resultado: dilo, no lo rellenes.

TEMA: {tema}

FICHAS (JSON):
{fichas}"""


# ---------------------------------------------------------------------------
# 5 · Contraste con PubMed (R2: solo PMIDs devueltos por la tool)
# ---------------------------------------------------------------------------
def _load_pubmed_tool():
    path = ROOT / "skills" / "pubmed_search" / "tool.py"
    spec = importlib.util.spec_from_file_location("skill_pubmed_search", path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.pubmed_search


def pubmed_contrast(terminos: list[str], retmax: int = 5) -> tuple[str, list[str]]:
    """Ejecuta la skill pubmed_search. Devuelve (texto crudo, PMIDs realmente devueltos)."""
    if not terminos:
        return "(sin términos de búsqueda)", []
    buscar = _load_pubmed_tool()
    bloques = []
    for term in terminos[:3]:               # 3 consultas distintas > 10 variantes de la misma
        out = buscar({"query": term, "retmax": retmax})
        bloques.append(f"### Consulta: {term}\n{out}")
    raw = "\n\n".join(bloques)
    return raw, sorted(set(re.findall(r"PMID\s+(\d{6,9})", raw)))


# ---------------------------------------------------------------------------
# Orquestación
# ---------------------------------------------------------------------------
def process_paper(path: Path, tema: str, guard: mp.CostGuard, *,
                  use_pubmed: bool = True, dry_run: bool = False) -> PaperRecord:
    rec = PaperRecord(archivo=path.name, formato=path.suffix.lower(), caracteres=0)
    try:
        raw = extract_text(path)
    except RuntimeError as e:
        rec.errores.append(str(e))
        return rec

    texto, redacciones = deidentify(raw)
    rec.caracteres, rec.redacciones_phi = len(texto), redacciones
    if len(texto) > PAPER_CHAR_BUDGET:
        texto, rec.truncado = texto[:PAPER_CHAR_BUDGET], True
    if not texto.strip():
        rec.errores.append("extracción vacía (¿PDF escaneado? requiere OCR)")
        return rec
    if dry_run:
        rec.ficha = {"(dry-run)": "extracción y de-identificación OK, sin llamadas al modelo"}
        return rec

    # Etapa 4 — ficha por paper: T2 (capacidad/velocidad por documento)
    try:
        salida = backends.one_shot(
            FICHA_PROMPT.format(tema=tema, archivo=path.name, texto=texto),
            task_class="synthesis", guard=guard, max_tokens=4000,
            system="Respondes solo con JSON válido. Nunca inventas datos ausentes del texto.")
        rec.ficha = _json_from(salida)
    except (ValueError, backends.BackendError, RuntimeError) as e:
        rec.errores.append(f"ficha: {e}")
        return rec

    # Etapa 5 — contraste con la literatura
    if not use_pubmed:
        rec.confianza = "media"
        rec.aporte = "(contraste con PubMed desactivado por --no-pubmed)"
        return rec
    try:
        pubmed_raw, pmids = pubmed_contrast(rec.ficha.get("terminos_pubmed") or [])
        rec.pmids_relacionados = pmids
        veredicto = _json_from(backends.one_shot(
            APORTE_PROMPT.format(ficha=json.dumps(rec.ficha, ensure_ascii=False, indent=1),
                                 pubmed=pubmed_raw),
            task_class="synthesis", guard=guard, max_tokens=3000,
            system="Respondes solo con JSON válido. Solo citas PMIDs presentes en el material."))
        citados = [p for p in veredicto.get("pmids_citados", []) if str(p) in pmids]
        rec.aporte = veredicto.get("aporte", "")
        rec.confianza = veredicto.get("confianza", "no evaluada")
        rec.ficha["novedad"] = veredicto.get("novedad", "no determinable")
        rec.ficha["no_verificado"] = veredicto.get("no_verificado", [])
        rec.ficha["pmids_citados"] = citados          # filtrados contra lo devuelto (R2)
    except (ValueError, backends.BackendError, RuntimeError) as e:
        rec.errores.append(f"contraste: {e}")
    return rec


def synthesize(records: list[PaperRecord], tema: str, guard: mp.CostGuard) -> str:
    """Etapa 6 — lectura transversal. Única llamada al tier profundo (T3)."""
    fichas = json.dumps([asdict(r) for r in records], ensure_ascii=False, indent=1)
    return backends.one_shot(
        SINTESIS_PROMPT.format(tema=tema, fichas=fichas),
        task_class="deep_analysis", guard=guard, max_tokens=16000,
        system="Eres un neurólogo académico. Rigor sobre fluidez. Declaras tus límites.")


def render_md(records: list[PaperRecord], sintesis: str, tema: str, guard: mp.CostGuard) -> str:
    filas = ["| Artículo | Diseño | n | Hallazgo principal | Novedad | Confianza |",
             "|---|---|---|---|---|---|"]
    for r in records:
        f = r.ficha or {}
        filas.append("| {} | {} | {} | {} | {} | {} |".format(
            (f.get("titulo") or r.archivo)[:70].replace("|", "/"),
            f.get("diseno") or "—", f.get("n") if f.get("n") is not None else "—",
            (f.get("hallazgo_principal") or "—")[:120].replace("|", "/"),
            f.get("novedad") or "—", r.confianza))

    detalle = []
    for r in records:
        f = r.ficha or {}
        pmids = ", ".join(f.get("pmids_citados") or []) or "(sin PMIDs verificados)"
        n_txt = f.get("n") if f.get("n") is not None else "—"
        detalle.append(
            f"### {f.get('titulo') or r.archivo}\n\n"
            f"- **Archivo:** `{r.archivo}`\n"
            f"- **Diseño / n:** {f.get('diseno') or '—'} / {n_txt}\n"
            f"- **Población:** {f.get('poblacion') or '—'}\n"
            f"- **Hallazgo:** {f.get('hallazgo_principal') or '—'}\n"
            f"- **Limitaciones:** {f.get('limitaciones') or '—'}\n"
            f"- **Calidad:** {f.get('calidad_metodologica') or '—'}\n"
            f"- **Aporte frente a la literatura:** {r.aporte or '—'}\n"
            f"- **PMIDs relacionados (devueltos por la tool):** {pmids}\n"
            + (f"- **Sin verificar:** {'; '.join(f.get('no_verificado') or [])}\n"
               if f.get("no_verificado") else "")
            + (f"- **Errores:** {'; '.join(r.errores)}\n" if r.errores else "")
            + ("- **Nota:** texto truncado al presupuesto de contexto; la ficha puede omitir "
               "material del final del artículo.\n" if r.truncado else ""))

    incidencias = [f"- `{r.archivo}`: {'; '.join(r.errores)}" for r in records if r.errores]
    return "\n".join([
        f"# Revisión: {tema}", "",
        f"*Generado el {date.today().isoformat()} · {len(records)} documentos analizados · "
        f"costo estimado de la corrida ${guard.spent:.4f} USD*", "",
        "## Tabla comparativa", "", *filas, "", sintesis, "",
        "## Fichas por artículo", "", *detalle, "",
        "## Incidencias", "", *(incidencias or ["- Ninguna."]), "",
        "---", "",
        "*Trazabilidad: los PMIDs listados provienen exclusivamente de consultas ejecutadas "
        "a PubMed en esta corrida; los datos de cada ficha provienen del texto del archivo "
        "correspondiente. Todo lo no anclado va marcado `[sin verificar]`.*", ""])


def run(folder: Path, out_dir: Path, tema: str, *, max_papers: int = 0,
        use_pubmed: bool = True, dry_run: bool = False) -> dict:
    papers = discover(folder)
    if max_papers:
        papers = papers[:max_papers]
    if not papers:
        raise FileNotFoundError(f"No hay PDF/DOCX/MD en {folder}")

    guard = mp.CostGuard()
    print(f"[paper_review] {len(papers)} documentos en {folder}")
    print(f"[paper_review] política: fichas → {mp.plan('synthesis').model} · "
          f"síntesis → {mp.plan('deep_analysis').model} · "
          f"techo ${guard.per_call:.2f}/llamada, ${guard.per_session:.2f}/corrida")

    records = []
    for i, p in enumerate(papers, 1):
        print(f"[{i}/{len(papers)}] {p.name}")
        records.append(process_paper(p, tema, guard, use_pubmed=use_pubmed, dry_run=dry_run))

    sintesis = ("*(dry-run: sin síntesis)*" if dry_run
                else synthesize(records, tema, guard))

    out_dir.mkdir(parents=True, exist_ok=True)
    payload = {"tema": tema, "fecha": date.today().isoformat(), "carpeta": str(folder),
               "costo_usd_estimado": guard.spent,
               "modelos": {"ficha": mp.plan("synthesis").model,
                           "sintesis": mp.plan("deep_analysis").model},
               "papers": [asdict(r) for r in records], "sintesis_md": sintesis}
    (out_dir / "revision.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    (out_dir / "revision.md").write_text(
        render_md(records, sintesis, tema, guard), encoding="utf-8")
    print(f"[paper_review] listo: {out_dir/'revision.md'} · {out_dir/'revision.json'} · "
          f"costo ${guard.spent:.4f}")
    return payload


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Análisis científico multi-paper con modelo por etapa.")
    ap.add_argument("--dir", required=True, help="Carpeta con los PDF/DOCX a analizar.")
    ap.add_argument("--out", default="revision", help="Carpeta de salida (revision.md + .json).")
    ap.add_argument("--tema", default="Revisión de literatura",
                    help="Pregunta o tema de la revisión.")
    ap.add_argument("--max-papers", type=int, default=0, help="Tope de documentos (0 = todos).")
    ap.add_argument("--no-pubmed", action="store_true", help="Salta el contraste con PubMed.")
    ap.add_argument("--dry-run", action="store_true",
                    help="Extrae y de-identifica sin llamar a ningún modelo.")
    a = ap.parse_args()
    try:
        run(Path(a.dir).expanduser(), Path(a.out).expanduser(), a.tema,
            max_papers=a.max_papers, use_pubmed=not a.no_pubmed, dry_run=a.dry_run)
    except (FileNotFoundError, RuntimeError) as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
