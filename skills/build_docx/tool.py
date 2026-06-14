"""Skill build_docx — materializa un documento Word desde secciones estructuradas."""
from __future__ import annotations
from pathlib import Path


def build_docx(args: dict) -> str:
    title = args.get("title", "Documento")
    sections = args.get("sections", [])
    out = args.get("out", "salida.docx")
    try:
        from docx import Document  # import perezoso
    except ImportError:
        return "ERROR: instala python-docx (pip install python-docx)."
    doc = Document()
    doc.add_heading(title, level=0)
    for sec in sections:
        if sec.get("heading"):
            doc.add_heading(str(sec["heading"]), level=1)
        for para in str(sec.get("body", "")).split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    Path(out).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return f"OK: documento creado en {out} ({len(sections)} secciones)."


def register_skill(reg) -> None:
    reg.register(
        "build_docx",
        "Crea un .docx desde {title, sections:[{heading, body}], out}.",
        {"type": "object",
         "properties": {"title": {"type": "string"},
                        "sections": {"type": "array"},
                        "out": {"type": "string"}},
         "required": ["title", "sections", "out"]},
        build_docx,
    )
